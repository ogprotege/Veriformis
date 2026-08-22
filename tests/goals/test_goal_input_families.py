"""Prove goal input-family eligibility against what each parser really recovers."""

from __future__ import annotations

import io
import re
from pathlib import Path

import pytest

from veriformis.goals import goal_catalog
from veriformis.ir import nodes
from veriformis.parsers.dispatch import parse_captured_source
from veriformis.rules.cleaning import plan_cleaning
from veriformis.rules.library import custom_regex
from veriformis.taxonomy import IMPLEMENTED_INPUT_FAMILIES, input_family_for_suffix

_PDF_SAMPLE = Path(__file__).parents[1] / "fixtures" / "group5" / "minimal-text.pdf"
_SYNTHETIC_PAGE_LABEL = re.compile(r"^Page \d+$")

_SUPPORTED_SCALARS = {
    "Heading": ("level",),
    "CodeBlock": ("language",),
    "Link": ("href", "title"),
    "Image": ("src", "title"),
    "Math": ("display",),
    "Citation": ("key", "locator"),
    "ListBlock": ("ordered",),
    "ListItem": ("checked",),
    "Table": ("alignments",),
}


def _docx_bytes() -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("Recovered heading", level=1)
    document.add_paragraph("Body text beneath the heading for recovery.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_SAMPLES: dict[str, tuple[str, bytes]] = {
    "plain-text": ("notes.txt", b"Plain paragraph one.\n\nPlain paragraph two.\n"),
    "source-code": ("tool.py", b"def run():\n    return 1\n"),
    "markdown": ("guide.md", b"# Recovered heading\n\nBody text [site](https://x.test).\n"),
    "word-document": ("memo.docx", b""),
    "html": ("page.html", b"<html><body><h1>Recovered heading</h1><p>Body text.</p></body></html>"),
    "delimited-table": ("rows.csv", b"name,value\nalpha,1\nbeta,2\n"),
    "json-records": ("records.jsonl", b'{"text": "alpha record"}\n{"text": "beta record"}\n'),
    "pdf-text": ("minimal-text.pdf", b""),
}


def _walk(block):
    yield block
    for child in getattr(block, "children", []) or []:
        if hasattr(child, "block_index"):
            yield from _walk(child)


def _heading_text(block) -> str:
    return "".join(getattr(child, "value", "") for child in block.children)


def _is_synthetic_page_label(block) -> bool:
    """PDF recovery emits per-page `Page N` headings; they are labels, not structure."""
    return isinstance(block, nodes.Heading) and bool(
        _SYNTHETIC_PAGE_LABEL.match(_heading_text(block))
    )


def _supplies(family: str) -> dict[str, bool]:
    name, raw = _SAMPLES[family]
    if family == "word-document":
        raw = _docx_bytes()
    if family == "pdf-text":
        raw = _PDF_SAMPLE.read_bytes()
    result = parse_captured_source(Path(name), logical_path=name, raw_bytes=raw)
    blocks = [block for top in result.document.children for block in _walk(top)]
    real_headings = [
        block
        for block in blocks
        if isinstance(block, nodes.Heading)
        and block.children
        and not _is_synthetic_page_label(block)
    ]
    body = any(isinstance(block, nodes.Paragraph) and block.children for block in blocks)
    scalar = False
    for block in blocks:
        if _is_synthetic_page_label(block):
            continue
        for field in _SUPPORTED_SCALARS.get(type(block).__name__, ()):
            if getattr(block, field, None) not in (None, "", [], ()):
                scalar = True
    # A recorded before/after pair needs cleaning to actually edit text.
    preview = plan_cleaning(result.document, [custom_regex("e", "3")])
    editable = sum(record.edits for record in preview.records) > 0
    return {
        "text": bool(blocks),
        "editable": editable,
        "heading": bool(real_headings) and body,
        "scalar": scalar,
    }


_EVIDENCE_KIND = {
    "learn-the-text": "text",
    "continue-a-passage": "text",
    "reproduce-a-recorded-change": "editable",
    "recover-a-section-from-its-heading": "heading",
    "extract-a-structured-value": "scalar",
}


@pytest.mark.parametrize("family", list(IMPLEMENTED_INPUT_FAMILIES))
def test_parsers_supply_exactly_the_evidence_each_goal_claims(family: str) -> None:
    name, _ = _SAMPLES[family]
    assert input_family_for_suffix(Path(name).suffix) == family
    supplies = _supplies(family)
    for goal in goal_catalog().goals:
        kind = _EVIDENCE_KIND[goal.goal_id]
        eligible = family in goal.eligible_input_families
        assert supplies[kind] == eligible, (goal.goal_id, family, supplies)


def test_source_code_is_never_edited_by_cleaning() -> None:
    """Binds the before/after exclusion to cleaning's non-editable code blocks."""
    name, raw = _SAMPLES["source-code"]
    result = parse_captured_source(Path(name), logical_path=name, raw_bytes=raw)
    assert all(isinstance(block, nodes.CodeBlock) for block in result.document.children)
    preview = plan_cleaning(result.document, [custom_regex("e", "3")])
    assert sum(record.edits for record in preview.records) == 0
    assert "source-code" not in goal_catalog().goal("reproduce-a-recorded-change").eligible_input_families


def test_pdf_text_headings_are_synthetic_page_labels() -> None:
    """Binds the PDF exclusions to the parser's per-page labels, not real headings."""
    result = parse_captured_source(
        _PDF_SAMPLE, logical_path="minimal-text.pdf", raw_bytes=_PDF_SAMPLE.read_bytes()
    )
    blocks = [block for top in result.document.children for block in _walk(top)]
    headings = [block for block in blocks if isinstance(block, nodes.Heading)]
    assert headings, "the PDF sample must exercise the heading path"
    assert all(_is_synthetic_page_label(block) for block in headings)
    assert any(isinstance(block, nodes.Paragraph) and block.children for block in blocks)
    catalog = goal_catalog()
    assert "pdf-text" not in catalog.goal("recover-a-section-from-its-heading").eligible_input_families
    assert "pdf-text" not in catalog.goal("extract-a-structured-value").eligible_input_families
    for goal_id in ("learn-the-text", "continue-a-passage", "reproduce-a-recorded-change"):
        assert "pdf-text" in catalog.goal(goal_id).eligible_input_families
