"""Adversarial compile-preflight coverage over the complete goal/family matrix.

This module deliberately owns the cross-product and defect-closure cases, while
``test_compile_preflight.py`` owns the basic response contract and surfaces.
Fixtures are raw sources so every assertion exercises the same parser, cleaning,
chunking, construction, and curation code used by a real compile.
"""

from __future__ import annotations

import hashlib
import io
from pathlib import Path

import pytest
import veriformis.sources as sources_module

from veriformis.errors import (
    ConstructionError,
    GoalCatalogError,
    InvalidSourceLocatorError,
)
from veriformis.goals import goal_catalog
from veriformis.goals.catalog import require_goal_input_family
from veriformis.goals.preflight import (
    MAX_RESPONSE_BYTES,
    MAX_SOURCE_BYTES,
    build_compile_preflight,
)
from veriformis.parsers.dispatch import parse_captured_source
from veriformis.pipeline import PipelineService
from veriformis.sources import derive_logical_paths
from veriformis.taxonomy import IMPLEMENTED_INPUT_FAMILIES


_PDF_SAMPLE = Path(__file__).parents[1] / "fixtures" / "group5" / "minimal-text.pdf"
_SERVICE = PipelineService()

_FAMILY_NAMES = {
    "plain-text": "notes.txt",
    "source-code": "tool.py",
    "markdown": "guide.md",
    "word-document": "memo.docx",
    "html": "page.html",
    "pdf-text": "minimal-text.pdf",
    "delimited-table": "rows.csv",
    "json-records": "records.jsonl",
}

_SOURCE_BYTES = {
    "plain-text": (
        b"First   source-grounded paragraph with enough text for a continuation.\n\n"
        b"Second paragraph remains exact source text.\n"
    ),
    "source-code": b"def source_grounded(value):\n    return value + 1\n",
    "markdown": (
        b"# Recovered heading\n\n"
        b"Body   text beneath the heading with exact source material.\n"
    ),
    "html": (
        b"<html><body><h1>Recovered heading</h1>"
        b"<p>Body   text beneath the heading with exact source material.</p>"
        b"</body></html>"
    ),
    "delimited-table": b"name,value\nalpha,first   value\nbeta,second value\n",
    "json-records": (
        b'{"text":"first   source-grounded record"}\n'
        b'{"text":"second source-grounded record"}\n'
    ),
}

_BEFORE_AFTER_CUSTOM = {
    "plain-text": "source-grounded",
    "source-code": "return",
    "markdown": "source material",
    "word-document": "source material",
    "html": "source material",
    "pdf-text": "Hello",
    "delimited-table": "alpha",
    "json-records": "first",
}

_EXPECTED_ELIGIBLE = {
    "learn-the-text": frozenset(IMPLEMENTED_INPUT_FAMILIES),
    "continue-a-passage": frozenset(IMPLEMENTED_INPUT_FAMILIES),
    "recover-a-section-from-its-heading": frozenset(
        {"markdown", "word-document", "html"}
    ),
    "reproduce-a-recorded-change": frozenset(
        set(IMPLEMENTED_INPUT_FAMILIES) - {"source-code"}
    ),
    "extract-a-structured-value": frozenset(
        {"source-code", "markdown", "word-document", "html"}
    ),
}

_MATRIX = tuple(
    pytest.param(
        goal_id,
        family,
        family in eligible,
        id=f"{goal_id}/{family}",
    )
    for goal_id, eligible in _EXPECTED_ELIGIBLE.items()
    for family in IMPLEMENTED_INPUT_FAMILIES
)


def _docx_bytes() -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("Recovered heading", level=1)
    document.add_paragraph(
        "Body   text beneath the heading with enough exact source material."
    )
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _raw_source(family: str) -> bytes:
    if family == "word-document":
        return _docx_bytes()
    if family == "pdf-text":
        return _PDF_SAMPLE.read_bytes()
    return _SOURCE_BYTES[family]


def _write_source(root: Path, family: str, *, stem: str = "source") -> Path:
    name = _FAMILY_NAMES[family]
    suffix = Path(name).suffix
    path = root / f"{stem}{suffix}"
    path.write_bytes(_raw_source(family))
    return path


def _tree_snapshot(root: Path) -> dict[str, str]:
    return {
        str(path.relative_to(root)): hashlib.sha256(path.read_bytes()).hexdigest()
        for path in sorted(root.rglob("*"))
        if path.is_file()
    }


def _compile_to_chunk(root: Path, family: str, goal_id: str) -> Path:
    source_root = root / "sources"
    source_root.mkdir()
    source = _write_source(source_root, family)
    workspace = root / "workspace"
    _SERVICE.parse([source], workspace, source_root=source_root)
    _SERVICE.clean(workspace)
    _SERVICE.chunk(workspace, goal=goal_id)
    return workspace


@pytest.mark.parametrize(("goal_id", "family", "expected"), _MATRIX)
def test_catalog_closes_the_exact_40_goal_by_family_cells(
    goal_id: str,
    family: str,
    expected: bool,
) -> None:
    """Freeze all 30 eligible and 10 refused cells before runtime assertions."""
    assert len(_MATRIX) == 40
    assert (family in goal_catalog().goal(goal_id).eligible_input_families) is expected


@pytest.mark.parametrize(("goal_id", "family", "expected"), _MATRIX)
def test_shared_runtime_family_gate_agrees_for_all_40_cells(
    goal_id: str,
    family: str,
    expected: bool,
) -> None:
    """The gate used by both preflight and construct must bind suffix to parser."""
    name = _FAMILY_NAMES[family]
    parsed = parse_captured_source(
        Path(name),
        logical_path=name,
        raw_bytes=_raw_source(family),
    )
    if expected:
        assert (
            require_goal_input_family(
                goal_id,
                logical_path=name,
                parser_id=parsed.source.parser,
            )
            == family
        )
        return
    with pytest.raises(
        GoalCatalogError,
        match=(rf"goal {goal_id!r} does not accept input family {family!r}"),
    ):
        require_goal_input_family(
            goal_id,
            logical_path=name,
            parser_id=parsed.source.parser,
        )


@pytest.mark.parametrize(("goal_id", "family", "expected"), _MATRIX)
def test_preflight_reports_the_exact_40_goal_by_family_verdicts_without_writes(
    tmp_path: Path,
    goal_id: str,
    family: str,
    expected: bool,
) -> None:
    source = _write_source(tmp_path, family)
    before = _tree_snapshot(tmp_path)

    report = build_compile_preflight(
        [source],
        source_root=tmp_path,
        goal=goal_id,
        # Keep this assertion about family and evidence. A separate test binds
        # the default evaluation-required split refusal for one leakage group.
        evaluation_required=False,
    )

    assert _tree_snapshot(tmp_path) == before
    assert report.counts.source_count == 1
    assert len(report.sources) == 1
    verdict = report.sources[0]
    assert verdict.logical_path == source.name
    assert verdict.input_family == family
    assert verdict.parser_id is not None
    assert verdict.parser_eligible is True
    assert verdict.goal_family_eligible is expected
    family_refusals = [
        item
        for item in verdict.refusal_reasons
        if item.code == "goal-input-family-ineligible"
    ]
    if expected:
        assert not family_refusals
        assert report.evaluated_through in {"construct", "curate", "split"}
    else:
        assert report.evaluated_through == "family"
        assert verdict.evidence_status == "not-evaluated"
        assert verdict.admitted is False
        assert len(family_refusals) == 1
        assert family_refusals[0].detail_codes == (family,)
        assert (
            f"goal {goal_id!r} does not accept input family {family!r}"
            in family_refusals[0].message
        )


@pytest.mark.parametrize(("goal_id", "family", "expected"), _MATRIX)
def test_preflight_and_real_stages_agree_for_all_40_cells(
    tmp_path: Path,
    goal_id: str,
    family: str,
    expected: bool,
) -> None:
    """Run the same raw source through preflight and persisted stage code."""
    source_root = tmp_path / "sources"
    source_root.mkdir()
    source = _write_source(source_root, family)
    # Before/after needs a recorded edit. The exact custom rule is applied to
    # both paths; it never supplies invented text and the code-family cell is
    # refused by the family gate before construction.
    custom = (
        _BEFORE_AFTER_CUSTOM[family] if goal_id == "reproduce-a-recorded-change" else ""
    )
    size = (
        (10 if family == "pdf-text" else 24)
        if goal_id == "reproduce-a-recorded-change"
        else None
    )
    overlap = 0 if size is not None else None
    report = build_compile_preflight(
        [source],
        source_root=source_root,
        goal=goal_id,
        custom=custom,
        size=size,
        overlap=overlap,
        evaluation_required=False,
    )

    workspace = tmp_path / "workspace"
    _SERVICE.parse([source], workspace, source_root=source_root)
    _SERVICE.clean(workspace, custom=custom)
    _SERVICE.chunk(workspace, goal=goal_id, size=size, overlap=overlap)
    before_construct = _tree_snapshot(workspace)

    if not expected:
        assert report.evaluated_through == "family"
        assert report.sources[0].goal_family_eligible is False
        with pytest.raises(ConstructionError, match="input-family admission failed"):
            _SERVICE.construct(workspace, goal=goal_id)
        assert _tree_snapshot(workspace) == before_construct
        return

    assert report.sources[0].goal_family_eligible is True
    assert report.sources[0].evidence_status == "available"
    assert report.admitted is True, (
        goal_id,
        family,
        report.missing_evidence,
        report.coverage_blockers,
        report.sources[0].refusal_reasons,
    )
    constructed = _SERVICE.construct(workspace, goal=goal_id)
    assert constructed.candidate_count == report.counts.candidate_count
    assert constructed.record_count == report.counts.record_count
    assert constructed.diagnostic_count == len(report.missing_evidence)
    curated = _SERVICE.curate(workspace, evaluation_required=False)
    assert curated.included_count == report.counts.included_count
    assert curated.excluded_count == report.counts.excluded_count
    assert curated.quarantined_count == report.counts.quarantined_count
    assert curated.coverage_blockers == ()
    split = _SERVICE.split(workspace)
    assert split.train_record_count == curated.included_count
    assert split.evaluation_record_count == 0


@pytest.mark.parametrize(
    "goal_id",
    ["recover-a-section-from-its-heading", "extract-a-structured-value"],
)
def test_pdf_synthetic_page_labels_never_pass_the_goal_family_gate(
    tmp_path: Path,
    goal_id: str,
) -> None:
    source = _write_source(tmp_path, "pdf-text")
    report = build_compile_preflight(
        [source], source_root=tmp_path, goal=goal_id, evaluation_required=False
    )
    verdict = report.sources[0]
    assert report.evaluated_through == "family"
    assert report.counts.candidate_count == report.counts.record_count == 0
    assert report.missing_evidence == ()
    assert verdict.goal_family_eligible is False
    assert [item.code for item in verdict.refusal_reasons] == [
        "goal-input-family-ineligible"
    ]


@pytest.mark.parametrize(
    ("goal_id", "name", "text", "diagnostic_code"),
    [
        (
            "learn-the-text",
            "empty.txt",
            "",
            "source-chunks-unavailable",
        ),
        (
            "continue-a-passage",
            "tiny.txt",
            "x",
            "continuation-boundary-unavailable",
        ),
        (
            "recover-a-section-from-its-heading",
            "flat.md",
            "Plain body text without a heading.",
            "section-structure-unavailable",
        ),
        (
            "reproduce-a-recorded-change",
            "unchanged.txt",
            "Ordinary source text with no cleaning edit.",
            "transformation-pair-unavailable",
        ),
        (
            "extract-a-structured-value",
            "flat.md",
            "Plain body text without a structured scalar.",
            "structured-field-unavailable",
        ),
    ],
)
def test_family_eligibility_does_not_overclaim_missing_goal_evidence(
    tmp_path: Path,
    goal_id: str,
    name: str,
    text: str,
    diagnostic_code: str,
) -> None:
    source = tmp_path / name
    source.write_text(text, encoding="utf-8")
    report = build_compile_preflight(
        [source], source_root=tmp_path, goal=goal_id, evaluation_required=False
    )
    verdict = report.sources[0]
    assert verdict.parser_eligible is True
    assert verdict.goal_family_eligible is True
    assert verdict.evidence_status == "missing"
    assert verdict.admitted is False
    assert report.counts.candidate_count == report.counts.record_count == 0
    assert [item.code for item in report.missing_evidence] == [diagnostic_code]
    evidence_refusal = next(
        item
        for item in verdict.refusal_reasons
        if item.code == "goal-evidence-unavailable"
    )
    assert evidence_refusal.detail_codes == (diagnostic_code,)
    assert report.coverage_blockers[0].blocker_codes == (
        "no-constructed-candidates",
        "no-dataset-records",
        "no-included-contribution",
    )


def test_expected_quality_exclusion_and_coverage_blocker_are_both_reported(
    tmp_path: Path,
) -> None:
    source = tmp_path / "passage.txt"
    source.write_text(
        "A sufficiently long exact source passage for a deterministic midpoint split.",
        encoding="utf-8",
    )
    report = build_compile_preflight(
        [source],
        source_root=tmp_path,
        goal="continue-a-passage",
        minimum_target_characters=100_000,
        evaluation_required=False,
    )
    assert report.counts.candidate_count == report.counts.record_count == 1
    assert report.counts.included_count == 0
    assert report.counts.excluded_count == 1
    assert [item.reason_codes for item in report.expected_exclusions] == [
        ("target-too-short",)
    ]
    assert [
        (item.stage, item.status, item.reason_code, item.count)
        for item in report.expected_exclusion_counts
    ] == [("curate", "excluded", "target-too-short", 1)]
    assert report.coverage_blockers[0].blocker_codes == ("no-included-contribution",)
    assert [reason.code for reason in report.sources[0].refusal_reasons] == [
        "curation-coverage-blocked"
    ]
    assert report.sources[0].refusal_reasons[0].detail_codes == (
        "no-included-contribution",
    )


def test_primary_source_cap_exclusion_preserves_complete_source_coverage(
    tmp_path: Path,
) -> None:
    source = tmp_path / "two-paragraphs.txt"
    source.write_text(
        "First independent paragraph.\n\nSecond independent paragraph.",
        encoding="utf-8",
    )
    report = build_compile_preflight(
        [source],
        source_root=tmp_path,
        goal="learn-the-text",
        size=20,
        overlap=0,
        balance_mode="primary-source-cap",
        maximum_records_per_primary_source=1,
        evaluation_required=False,
    )
    assert report.admitted is True
    assert report.counts.record_count == 2
    assert report.counts.included_count == 1
    assert report.counts.excluded_count == 1
    assert [item.reason_codes for item in report.expected_exclusions] == [
        ("primary-source-cap",)
    ]
    assert report.coverage_blockers == ()


def test_curation_deduplicates_globally_across_sources_before_admission(
    tmp_path: Path,
) -> None:
    text = "The same sufficiently long source-grounded continuation passage."
    paths = []
    for name in ("a.txt", "b.txt"):
        path = tmp_path / name
        path.write_text(text, encoding="utf-8")
        paths.append(path)
    report = build_compile_preflight(
        paths,
        source_root=tmp_path,
        goal="continue-a-passage",
        evaluation_required=False,
    )
    assert report.counts.record_count == 2
    assert report.counts.included_count == 1
    assert report.counts.excluded_count == 1
    assert [item.reason_codes for item in report.expected_exclusions] == [
        ("exact-duplicate",)
    ]
    assert len(report.coverage_blockers) == 1
    assert report.coverage_blockers[0].blocker_codes == ("no-included-contribution",)
    assert sum(source.admitted for source in report.sources) == 1


def test_curation_quarantines_conflicting_structured_targets_from_one_context(
    tmp_path: Path,
) -> None:
    source = tmp_path / "links.md"
    source.write_text(
        "[one](https://one.test) and [two](https://two.test)",
        encoding="utf-8",
    )
    report = build_compile_preflight(
        [source],
        source_root=tmp_path,
        goal="extract-a-structured-value",
        evaluation_required=False,
    )
    assert report.counts.record_count == 2
    assert report.counts.included_count == 0
    assert report.counts.quarantined_count == 2
    assert [item.reason_codes for item in report.expected_exclusions] == [
        ("conflicting-target",),
        ("conflicting-target",),
    ]
    assert [
        (item.stage, item.status, item.reason_code, item.count)
        for item in report.expected_exclusion_counts
    ] == [("curate", "quarantined", "conflicting-target", 2)]


def test_default_required_evaluation_predicts_the_single_group_split_refusal(
    tmp_path: Path,
) -> None:
    source = tmp_path / "one.txt"
    source.write_text(
        "One complete source-grounded training passage.", encoding="utf-8"
    )
    report = build_compile_preflight(
        [source], source_root=tmp_path, goal="learn-the-text"
    )
    assert report.coverage_blockers == ()
    assert report.counts.included_count == 1
    assert report.evaluated_through == "curate"
    assert report.admitted is False
    split_refusals = [
        reason
        for reason in report.sources[0].refusal_reasons
        if reason.code == "evaluation-partition-unavailable"
    ]
    assert len(split_refusals) == 1
    assert split_refusals[0].detail_codes == ("split-invalid",)


def test_two_independent_groups_satisfy_the_default_required_evaluation_split(
    tmp_path: Path,
) -> None:
    paths = []
    for name, text in (
        ("a.txt", "First complete source-grounded training passage."),
        ("b.txt", "Second independent source-grounded training passage."),
    ):
        path = tmp_path / name
        path.write_text(text, encoding="utf-8")
        paths.append(path)
    report = build_compile_preflight(paths, source_root=tmp_path, goal="learn-the-text")
    assert report.evaluated_through == "split"
    assert report.admitted is True
    assert report.counts.admitted_source_count == 2
    assert all(not source.refusal_reasons for source in report.sources)


@pytest.mark.parametrize(
    "goal_id",
    ["recover-a-section-from-its-heading", "extract-a-structured-value"],
)
def test_real_construct_shares_the_pdf_family_refusal_without_committing(
    tmp_path: Path,
    goal_id: str,
) -> None:
    workspace = _compile_to_chunk(tmp_path, "pdf-text", goal_id)
    before = _tree_snapshot(workspace)
    with pytest.raises(
        ConstructionError,
        match=(
            rf"goal input-family admission failed: goal {goal_id!r} does not accept "
            r"input family 'pdf-text'"
        ),
    ):
        _SERVICE.construct(workspace, goal=goal_id)
    assert _tree_snapshot(workspace) == before


def test_duplicate_source_paths_fail_before_capture_or_workspace_creation(
    tmp_path: Path,
) -> None:
    source = tmp_path / "same.txt"
    source.write_text("same source", encoding="utf-8")
    with pytest.raises(InvalidSourceLocatorError, match="resolve to the same file"):
        derive_logical_paths([source, source], source_root=tmp_path)

    report = build_compile_preflight(
        [source, source],
        source_root=tmp_path,
        goal="learn-the-text",
        evaluation_required=False,
    )
    assert report.evaluated_through == "capture"
    assert report.captured_source_digest is None
    assert report.counts.source_count == 2
    assert all(
        [reason.code for reason in verdict.refusal_reasons] == ["source-read-failed"]
        for verdict in report.sources
    )
    workspace = tmp_path / "workspace"
    with pytest.raises(InvalidSourceLocatorError, match="resolve to the same file"):
        _SERVICE.parse([source, source], workspace, source_root=tmp_path)
    assert not workspace.exists()


def test_shared_runtime_family_gate_rejects_suffix_parser_disagreement() -> None:
    with pytest.raises(GoalCatalogError, match="parser 'markdown'.*not one of"):
        require_goal_input_family(
            "learn-the-text",
            logical_path="notes.txt",
            parser_id="markdown",
        )


def test_symlink_escape_is_refused_before_capture_or_workspace_creation(
    tmp_path: Path,
) -> None:
    source_root = tmp_path / "root"
    source_root.mkdir()
    outside = tmp_path / "outside.txt"
    outside.write_text("outside source", encoding="utf-8")
    link = source_root / "escape.txt"
    link.symlink_to(outside)

    report = build_compile_preflight(
        [link],
        source_root=source_root,
        goal="learn-the-text",
        evaluation_required=False,
    )
    assert report.evaluated_through == "capture"
    assert report.captured_source_digest is None
    assert [reason.code for reason in report.sources[0].refusal_reasons] == [
        "source-read-failed"
    ]
    assert (
        "source symlinks are not allowed"
        in report.sources[0].refusal_reasons[0].message
    )
    workspace = tmp_path / "workspace"
    with pytest.raises(
        InvalidSourceLocatorError, match="source symlinks are not allowed"
    ):
        _SERVICE.parse([link], workspace, source_root=source_root)
    assert not workspace.exists()


def test_parent_directory_symlink_is_refused_by_preflight_and_parse(
    tmp_path: Path,
) -> None:
    source_root = tmp_path / "root"
    real_directory = source_root / "real"
    real_directory.mkdir(parents=True)
    (real_directory / "source.txt").write_text("inside source", encoding="utf-8")
    alias = source_root / "alias"
    alias.symlink_to(real_directory, target_is_directory=True)
    source = alias / "source.txt"

    report = build_compile_preflight(
        [source],
        source_root=source_root,
        goal="learn-the-text",
        evaluation_required=False,
    )
    assert report.evaluated_through == "capture"
    assert [reason.code for reason in report.sources[0].refusal_reasons] == [
        "source-read-failed"
    ]
    assert (
        "source symlinks are not allowed"
        in report.sources[0].refusal_reasons[0].message
    )
    workspace = tmp_path / "workspace"
    with pytest.raises(
        InvalidSourceLocatorError, match="source symlinks are not allowed"
    ):
        _SERVICE.parse([source], workspace, source_root=source_root)
    assert not workspace.exists()


def test_parent_directory_swap_to_symlink_is_refused_by_preflight_and_parse(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source_root = tmp_path / "root"
    parent = source_root / "parent"
    parent.mkdir(parents=True)
    source = parent / "source.txt"
    source.write_text("inside bytes", encoding="utf-8")
    outside_parent = tmp_path / "outside"
    outside_parent.mkdir()
    (outside_parent / "source.txt").write_text(
        "outside bytes must never be captured",
        encoding="utf-8",
    )
    saved_parent = source_root / "saved-parent"
    shared_capture = sources_module._capture_at
    preflight_swapped = False

    def swap_parent_before_preflight_capture(
        root_descriptor: int,
        relative: Path,
        path: Path,
    ) -> bytes:
        nonlocal preflight_swapped
        if not preflight_swapped:
            preflight_swapped = True
            parent.rename(saved_parent)
            parent.symlink_to(outside_parent, target_is_directory=True)
        return shared_capture(root_descriptor, relative, path)

    monkeypatch.setattr(
        sources_module,
        "_capture_at",
        swap_parent_before_preflight_capture,
    )
    report = build_compile_preflight(
        [source],
        source_root=source_root,
        goal="learn-the-text",
        evaluation_required=False,
    )
    assert preflight_swapped is True
    assert report.evaluated_through == "capture"
    assert report.sources[0].sha256 is None
    assert (
        "source symlinks are not allowed"
        in report.sources[0].refusal_reasons[0].message
    )

    monkeypatch.setattr(sources_module, "_capture_at", shared_capture)
    parent.unlink()
    saved_parent.rename(parent)
    parse_swapped = False

    def swap_parent_before_parse_capture(
        root_descriptor: int,
        relative: Path,
        path: Path,
    ) -> bytes:
        nonlocal parse_swapped
        if not parse_swapped:
            parse_swapped = True
            parent.rename(saved_parent)
            parent.symlink_to(outside_parent, target_is_directory=True)
        return shared_capture(root_descriptor, relative, path)

    monkeypatch.setattr(
        sources_module,
        "_capture_at",
        swap_parent_before_parse_capture,
    )
    workspace = tmp_path / "workspace"
    with pytest.raises(
        InvalidSourceLocatorError, match="source symlinks are not allowed"
    ):
        _SERVICE.parse([source], workspace, source_root=source_root)
    assert parse_swapped is True
    assert not workspace.exists()


def test_final_component_symlink_swap_is_refused_by_preflight_and_parse(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source_root = tmp_path / "root"
    source_root.mkdir()
    source = source_root / "source.txt"
    outside = tmp_path / "outside.txt"
    outside.write_text("outside bytes must never be captured", encoding="utf-8")

    source.write_text("inside bytes", encoding="utf-8")
    shared_capture = sources_module._capture_at
    preflight_swapped = False

    def swap_before_preflight_capture(
        root_descriptor: int,
        relative: Path,
        logical_path: str,
    ) -> tuple[bytes, tuple[int, int]]:
        nonlocal preflight_swapped
        if logical_path == "source.txt" and not preflight_swapped:
            preflight_swapped = True
            source.unlink()
            source.symlink_to(outside)
        return shared_capture(root_descriptor, relative, logical_path)

    monkeypatch.setattr(
        sources_module,
        "_capture_at",
        swap_before_preflight_capture,
    )
    report = build_compile_preflight(
        [source],
        source_root=source_root,
        goal="learn-the-text",
        evaluation_required=False,
    )
    assert preflight_swapped is True
    assert report.evaluated_through == "capture"
    assert report.sources[0].sha256 is None
    assert (
        "source symlinks are not allowed"
        in report.sources[0].refusal_reasons[0].message
    )

    monkeypatch.setattr(sources_module, "_capture_at", shared_capture)
    source.unlink()
    source.write_text("inside bytes", encoding="utf-8")
    parse_swapped = False

    def swap_before_parse_capture(
        root_descriptor: int,
        relative: Path,
        logical_path: str,
    ) -> tuple[bytes, tuple[int, int]]:
        nonlocal parse_swapped
        if logical_path == "source.txt" and not parse_swapped:
            parse_swapped = True
            source.unlink()
            source.symlink_to(outside)
        return shared_capture(root_descriptor, relative, logical_path)

    monkeypatch.setattr(
        sources_module,
        "_capture_at",
        swap_before_parse_capture,
    )
    workspace = tmp_path / "workspace"
    with pytest.raises(
        InvalidSourceLocatorError, match="source symlinks are not allowed"
    ):
        _SERVICE.parse([source], workspace, source_root=source_root)
    assert parse_swapped is True
    assert not workspace.exists()


@pytest.mark.parametrize(
    ("name", "raw", "evaluated_through", "refusal_code", "detail_code"),
    [
        (
            "unknown.xyz",
            b"unsupported suffix",
            "parse",
            "unsupported-input",
            "taxonomy-invalid",
        ),
        (
            "broken.jsonl",
            b"{not-json}\n",
            "parse",
            "parser-refused",
            "jsonl.invalid-line",
        ),
        (
            "missing.txt",
            None,
            "capture",
            "source-read-failed",
            "FileNotFoundError",
        ),
    ],
)
def test_capture_and_parser_refusals_keep_stable_codes(
    tmp_path: Path,
    name: str,
    raw: bytes | None,
    evaluated_through: str,
    refusal_code: str,
    detail_code: str,
) -> None:
    source = tmp_path / name
    if raw is not None:
        source.write_bytes(raw)
    report = build_compile_preflight(
        [source],
        source_root=tmp_path,
        goal="learn-the-text",
        evaluation_required=False,
    )
    verdict = report.sources[0]
    assert report.evaluated_through == evaluated_through
    assert verdict.admitted is False
    assert [
        (reason.code, reason.detail_codes) for reason in verdict.refusal_reasons
    ] == [(refusal_code, (detail_code,))]


def test_preflight_reads_each_source_once_and_binds_the_old_bytes_during_a_race(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "racing.txt"
    captured = b"Original source-grounded training passage."
    replacement = b"Mutated after the immutable capture."
    source.write_bytes(captured)
    original_capture = sources_module._capture_at
    source_reads = 0

    def racing_capture(
        root_descriptor: int,
        relative: Path,
        logical_path: str,
    ) -> tuple[bytes, tuple[int, int]]:
        nonlocal source_reads
        raw = original_capture(root_descriptor, relative, logical_path)
        if logical_path == "racing.txt":
            source_reads += 1
            if source_reads == 1:
                source.write_bytes(replacement)
        return raw

    monkeypatch.setattr(sources_module, "_capture_at", racing_capture)
    report = build_compile_preflight(
        [source],
        source_root=tmp_path,
        goal="learn-the-text",
        evaluation_required=False,
    )
    verdict = report.sources[0]
    assert source_reads == 1
    assert source.read_bytes() == replacement
    assert verdict.sha256 == hashlib.sha256(captured).hexdigest()
    assert verdict.size == len(captured)
    assert verdict.admitted is True
    assert "point-in-time-source-capture" in {
        item.code for item in report.known_limitations
    }


def test_unicode_transport_is_ascii_safe_and_within_public_bounds(
    tmp_path: Path,
) -> None:
    source = tmp_path / "unicode.txt"
    source.write_text(
        "Café — naïve source text with 漢字 and emoji U0001f642.",
        encoding="utf-8",
    )
    report = build_compile_preflight(
        [source],
        source_root=tmp_path,
        goal="learn-the-text",
        evaluation_required=False,
    )
    transport = report.transport_text().encode("utf-8")
    assert transport.isascii()
    assert len(transport) <= MAX_RESPONSE_BYTES
    verdict_transport = report.sources[0].model_dump_json().encode("utf-8")
    assert len(verdict_transport) <= MAX_SOURCE_BYTES


def test_many_parser_diagnostics_are_omitted_whole_to_preserve_bounds(
    tmp_path: Path,
) -> None:
    source = tmp_path / "many-unused-footnotes.md"
    source.write_text(
        "\n".join(f"[^{index}]: unused note {index}" for index in range(400)),
        encoding="utf-8",
    )
    report = build_compile_preflight(
        [source],
        source_root=tmp_path,
        goal="learn-the-text",
        evaluation_required=False,
    )
    verdict = report.sources[0]
    assert report.evaluated_through == "parse"
    assert verdict.omission_reason == "exact-source-exceeds-preflight-limit"
    assert verdict.diagnostics == ()
    assert verdict.omitted_diagnostic_count == 400
    assert len(report.transport_text().encode("utf-8")) <= MAX_RESPONSE_BYTES
    assert len(verdict.model_dump_json().encode("utf-8")) <= MAX_SOURCE_BYTES
    assert {reason.code for reason in verdict.refusal_reasons} == {"parser-refused"}
    assert {
        code for reason in verdict.refusal_reasons for code in reason.detail_codes
    } == {"markdown.unused-footnote-definition-refused"}


def test_total_response_budget_omits_each_source_diagnostic_set_whole(
    tmp_path: Path,
) -> None:
    paths = []
    for source_index in range(20):
        source = tmp_path / f"source-{source_index:02d}.md"
        source.write_text(
            "\n".join(
                f"[^{source_index}-{diagnostic_index}]: unused note"
                for diagnostic_index in range(30)
            ),
            encoding="utf-8",
        )
        paths.append(source)
    report = build_compile_preflight(
        paths,
        source_root=tmp_path,
        goal="learn-the-text",
        evaluation_required=False,
    )
    assert len(report.transport_text().encode("utf-8")) <= MAX_RESPONSE_BYTES
    assert len(report.sources) == 20
    assert all(source.diagnostics == () for source in report.sources)
    assert all(source.omitted_diagnostic_count == 30 for source in report.sources)
    assert {source.omission_reason for source in report.sources} == {
        "exact-source-exceeds-response-budget"
    }
