import copy
import hashlib
from io import BytesIO

import pytest
from docx import Document as DocxDocument

from veriformis.cli import _load_sources
from veriformis.diagnostics import (
    DiagnosticLocation,
    make_diagnostic,
    make_parse_report,
    parse_report_from_dict,
    parse_report_to_dict,
    validate_parse_report_locations,
)
from veriformis.errors import EvidenceError, ParseError, WorkspaceCorruptError
from veriformis.identity import derive_source_id, lossless_json_bytes
from veriformis.ir import Document, Paragraph, Span, Text, document_to_dict
from veriformis.workspace import (
    SourceDescriptor,
    Workspace,
    WorkspaceTransaction,
)


def _diagnostic(source_id: str, location: DiagnosticLocation):
    return make_diagnostic(
        source_id=source_id,
        parser_name="text",
        parser_version="1",
        code="test.location",
        severity="warning",
        disposition="normalized",
        loss_kind="presentation",
        location=location,
        message="Located parser observation.",
    )


def _located_report_value() -> dict:
    raw_sha = hashlib.sha256(b"one line").hexdigest()
    source_id = derive_source_id("source.txt", raw_sha)
    report = make_parse_report(
        source_id=source_id,
        parser_name="text",
        parser_version="1",
        diagnostics=[
            _diagnostic(
                source_id,
                DiagnosticLocation(kind="text", line_start=1, line_end=1),
            )
        ],
    )
    return parse_report_to_dict(report)


@pytest.mark.parametrize(
    ("updates", "message"),
    [
        (
            {"line_start": None, "line_end": None},
            "requires a line or raw-byte range",
        ),
        ({"line_end": None}, "requires both line_start and line_end"),
        (
            {"line_start": 2, "line_end": 1},
            "line_start/line_end is backward or invalid",
        ),
        (
            {
                "line_start": None,
                "line_end": None,
                "raw_byte_start": 0,
                "raw_byte_end": None,
            },
            "requires both raw_byte_start and raw_byte_end",
        ),
        (
            {
                "line_start": None,
                "line_end": None,
                "column_start": 1,
                "column_end": 2,
                "raw_byte_start": 0,
                "raw_byte_end": 1,
            },
            "columns require a line range",
        ),
        ({"part": "word/document.xml"}, "cannot carry OOXML coordinates"),
        (
            {
                "kind": "ooxml",
                "line_start": None,
                "line_end": None,
                "part": "",
                "xpath": "/",
            },
            "requires nonempty part and xpath",
        ),
        (
            {
                "kind": "ooxml",
                "part": "word/document.xml",
                "xpath": "/w:document",
            },
            "cannot carry text or raw-byte coordinates",
        ),
        (
            {"kind": "source"},
            "cannot carry text or raw-byte coordinates",
        ),
    ],
)
def test_parse_report_loader_rejects_malformed_location_shapes(updates, message):
    value = copy.deepcopy(_located_report_value())
    value["diagnostics"][0]["location"].update(updates)

    with pytest.raises(ParseError, match=message):
        parse_report_from_dict(value)


def test_location_creation_rejects_incomplete_and_backward_pairs():
    with pytest.raises(ValueError, match="requires both line_start and line_end"):
        DiagnosticLocation(kind="text", line_start=1)
    with pytest.raises(ValueError, match="backward or invalid"):
        DiagnosticLocation(kind="text", raw_byte_start=2, raw_byte_end=1)
    with pytest.raises(ValueError, match="must be integers"):
        DiagnosticLocation(kind="text", line_start=True, line_end=1)


def test_source_location_is_a_whole_source_sentinel():
    raw_sha = hashlib.sha256(b"opaque").hexdigest()
    source_id = derive_source_id("opaque.bin", raw_sha)
    report = make_parse_report(
        source_id=source_id,
        parser_name="text",
        parser_version="1",
        diagnostics=[_diagnostic(source_id, DiagnosticLocation(kind="source"))],
    )

    loaded = parse_report_from_dict(parse_report_to_dict(report))

    assert loaded.diagnostics[0].location == DiagnosticLocation(kind="source")


def test_parse_report_rejects_duplicate_diagnostic_identities():
    value = _located_report_value()
    value["diagnostics"].append(copy.deepcopy(value["diagnostics"][0]))

    with pytest.raises(ParseError, match="duplicate diagnostic identities"):
        parse_report_from_dict(value)


def _commit_parse_with_location(
    workspace: Workspace,
    location: DiagnosticLocation,
):
    raw = b"only line"
    logical_path = "source.txt"
    raw_sha = hashlib.sha256(raw).hexdigest()
    source_id = derive_source_id(logical_path, raw_sha)
    parser_config = {
        "parser": "text",
        "parser_version": "1",
        "canonical_stream_contract_version": 1,
    }
    with workspace.begin("parse") as transaction:
        raw_artifact = transaction.put_artifact(
            raw,
            kind="raw-source",
            media_type="application/octet-stream",
            source_ids=(source_id,),
            producer_id="veriformis.source-capture",
            producer_version="1",
            config={"logical_path": logical_path},
        )
        canonical_artifact = transaction.put_artifact(
            raw,
            kind="canonical-source-text",
            media_type="text/plain; charset=utf-8",
            source_ids=(source_id,),
            producer_id="veriformis.parser.text",
            producer_version="1",
            config=parser_config,
        )
        document_artifact = transaction.put_artifact(
            lossless_json_bytes(
                document_to_dict(
                    Document(
                        children=[
                            Paragraph(
                                children=[Text("only line")],
                                span=Span(0, len("only line")),
                                block_index=0,
                            )
                        ],
                        source_id=source_id,
                    )
                )
            ),
            kind="document-ir",
            media_type="application/json",
            source_ids=(source_id,),
            producer_id="veriformis.parser.text",
            producer_version="1",
            config=parser_config,
        )
        report = make_parse_report(
            source_id=source_id,
            parser_name="text",
            parser_version="1",
            diagnostics=[_diagnostic(source_id, location)],
        )
        diagnostics_artifact = transaction.put_artifact(
            lossless_json_bytes(parse_report_to_dict(report)),
            kind="parse-report",
            media_type="application/json",
            source_ids=(source_id,),
            producer_id="veriformis.parser.text",
            producer_version="1",
            config=parser_config,
        )
        source = SourceDescriptor.create(
            logical_path=logical_path,
            sha256=raw_sha,
            size=len(raw),
            parser_id="text",
            parser_version="1",
            raw_artifact_id=raw_artifact.id,
            extracted_artifact_id=canonical_artifact.id,
            document_artifact_id=document_artifact.id,
        )
        transaction.set_sources((source,))
        registry = transaction.put_artifact(
            lossless_json_bytes(
                [source.model_dump(mode="json", exclude={"original_path"})]
            ),
            kind="source-registry",
            media_type="application/json",
            source_ids=(source_id,),
            producer_id="veriformis.parse-stage",
            producer_version="1",
            config={"source_count": 1},
        )
        return transaction.commit(
            outputs={
                "registry": registry,
                f"source/{source_id}/raw": raw_artifact,
                f"source/{source_id}/canonical": canonical_artifact,
                f"source/{source_id}/document": document_artifact,
                f"source/{source_id}/diagnostics": diagnostics_artifact,
            },
            config={"sources": [logical_path]},
        )


@pytest.mark.parametrize(
    "location",
    [
        DiagnosticLocation(kind="text", line_start=2, line_end=2),
        DiagnosticLocation(kind="text", raw_byte_start=0, raw_byte_end=10),
        DiagnosticLocation(
            kind="text",
            line_start=1,
            line_end=1,
            column_start=1,
            column_end=11,
        ),
    ],
)
def test_parse_transaction_rejects_locations_outside_captured_input(
    tmp_path,
    location,
):
    workspace = Workspace.create(tmp_path / "ws")
    before = workspace.head_id

    with pytest.raises(WorkspaceCorruptError, match="locations do not match"):
        _commit_parse_with_location(workspace, location)

    assert workspace.head_id == before


def test_cli_loader_rechecks_locations_against_captured_input(tmp_path, monkeypatch):
    workspace = Workspace.create(tmp_path / "ws")
    monkeypatch.setattr(
        WorkspaceTransaction,
        "_validate_stage_semantics",
        lambda self, revision: None,
    )
    revision = _commit_parse_with_location(
        workspace,
        DiagnosticLocation(kind="text", line_start=2, line_end=2),
    )

    with pytest.raises(EvidenceError, match="exceed the captured source"):
        _load_sources(workspace, revision)


@pytest.mark.parametrize(
    ("part", "xpath", "message"),
    (
        (
            "word/does-not-exist.xml",
            "/w:document",
            "part does not exist exactly once",
        ),
        (
            "word/document.xml",
            "/w:document/w:body/w:p[999]",
            "xpath does not resolve",
        ),
    ),
)
def test_ooxml_locations_must_resolve_in_captured_docx(part, xpath, message):
    document = DocxDocument()
    document.add_paragraph("Body")
    stream = BytesIO()
    document.save(stream)
    raw = stream.getvalue()
    source_id = derive_source_id(
        "source.docx",
        hashlib.sha256(raw).hexdigest(),
    )
    diagnostic = make_diagnostic(
        source_id=source_id,
        parser_name="docx",
        parser_version="1",
        code="test.ooxml-location",
        severity="warning",
        disposition="normalized",
        loss_kind="presentation",
        location=DiagnosticLocation(kind="ooxml", part=part, xpath=xpath),
        message="Located DOCX observation.",
    )
    report = make_parse_report(
        source_id=source_id,
        parser_name="docx",
        parser_version="1",
        diagnostics=[diagnostic],
    )

    with pytest.raises(ParseError, match=message):
        validate_parse_report_locations(report, raw)
