import hashlib
from pathlib import Path

import pytest
from docx import Document as DocxBuilder
from markdown_it.token import Token

from veriformis.contracts import CANONICAL_STREAM_CONTRACT_VERSION
from veriformis.diagnostics import parse_report_from_dict, parse_report_to_dict
from veriformis.errors import ParseError
from veriformis.parsers.docx import parse_docx_file
from veriformis.parsers.markdown import parse_md_file
from veriformis.parsers.text import parse_text


def test_each_supported_parser_returns_versioned_diagnostics(tmp_path):
    text = tmp_path / "plain.txt"
    text.write_text("first\n\nsecond", encoding="utf-8")
    code = tmp_path / "sample.py"
    code.write_text("print('ok')\n", encoding="utf-8")
    markdown = tmp_path / "sample.md"
    markdown.write_text("# Title\n\nbody", encoding="utf-8")
    docx = tmp_path / "sample.docx"
    builder = DocxBuilder()
    builder.add_paragraph("Body")
    builder.save(docx)

    results = [
        parse_text(text, logical_path=text.name),
        parse_text(code, language="python", logical_path=code.name),
        parse_md_file(markdown, logical_path=markdown.name),
        parse_docx_file(docx, logical_path=docx.name),
    ]
    for result in results:
        report = result.diagnostics
        assert report is not None
        assert report.schema_version == "veriformis.parse-report/v1"
        assert report.source_id == result.source.id
        assert len(report.report_digest) == 64
        assert result.source.artifact_id.startswith("art-v1-")
        assert (
            result.source.canonical_stream_contract_version
            == CANONICAL_STREAM_CONTRACT_VERSION
        )
        assert len(result.source.stream_sha256) == 64
        assert result.document.source_id == result.source.id
        assert (
            hashlib.sha256(result.source.extracted_text.encode("utf-8")).hexdigest()
            == result.source.stream_sha256
        )


def test_parse_report_persistence_is_strict_and_identity_checked(tmp_path):
    path = tmp_path / "report.txt"
    path.write_text("Body", encoding="utf-8")
    report = parse_text(path, logical_path=path.name).diagnostics
    value = parse_report_to_dict(report)

    assert parse_report_from_dict(value) == report

    value["report_digest"] = "0" * 64
    with pytest.raises(ParseError, match="digest mismatch"):
        parse_report_from_dict(value)


def test_unsupported_markdown_construct_is_never_silent(tmp_path):
    path = tmp_path / "html.md"
    path.write_text(
        "before\n\n<section><b>lost structure</b></section>\n\nafter", encoding="utf-8"
    )

    first = parse_md_file(path, logical_path=path.name)
    second = parse_md_file(path, logical_path=path.name)
    assert first.diagnostics is not None
    diagnostics = first.diagnostics.diagnostics
    assert diagnostics
    html = [item for item in diagnostics if item.code.startswith("markdown.html-")]
    assert html
    assert all(item.severity == "warning" for item in html)
    assert all(item.disposition == "omitted" for item in html)
    assert all(
        item.location.kind == "text" and item.location.line_start for item in html
    )
    assert any(
        item.code == "markdown.html-block-omitted" and item.loss_kind == "text"
        for item in html
    )
    assert [item.diagnostic_id for item in html] == [
        item.diagnostic_id
        for item in second.diagnostics.diagnostics
        if item.code.startswith("markdown.html-")
    ]


def test_pandoc_metadata_removal_is_located_and_never_silent(tmp_path):
    path = tmp_path / "pandoc.md"
    path.write_text(
        "# Heading {#heading .major}\n\nBody []{#anchor .target} text",
        encoding="utf-8",
    )

    result = parse_md_file(path, logical_path=path.name)

    codes = [item.code for item in result.diagnostics.diagnostics]
    assert "markdown.pandoc-attributes-omitted" in codes
    assert "markdown.pandoc-anchor-omitted" in codes
    matching = [
        item
        for item in result.diagnostics.diagnostics
        if item.code.startswith("markdown.pandoc-")
    ]
    assert all(item.disposition == "omitted" for item in matching)
    assert all(item.loss_kind == "metadata" for item in matching)
    assert all(item.location.line_start is not None for item in matching)


def test_unknown_markdown_tokens_are_located_and_never_silent(tmp_path, monkeypatch):
    mystery = Token("mystery_block", "mystery", 0)
    mystery.map = [0, 1]
    mystery.content = "unrepresented"

    class Parser:
        def parse(self, _source, _environment=None):
            return [mystery]

    monkeypatch.setattr("veriformis.parsers.markdown._make_parser", lambda: Parser())
    result = parse_md_file(
        tmp_path / "mystery.md",
        raw_bytes=b"unrepresented",
        logical_path="mystery.md",
    )

    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "markdown.unsupported-token-omitted"
    )
    assert item.disposition == "omitted"
    assert item.location.line_start == 1
    assert item.details["token_type"] == "mystery_block"


def test_docx_reports_unavailable_page_provenance(tmp_path):
    path = tmp_path / "page.docx"
    builder = DocxBuilder()
    builder.add_paragraph("Body")
    builder.save(path)

    result = parse_docx_file(path, logical_path=path.name)
    assert result.diagnostics is not None
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "docx.page-provenance-unavailable"
    )
    assert item.location.kind == "ooxml"
    assert item.location.part == "word/document.xml"


def test_parsers_use_supplied_bytes_and_logical_path_without_rereading(
    tmp_path, monkeypatch
):
    docx_path = tmp_path / "captured.docx"
    builder = DocxBuilder()
    builder.add_paragraph("Captured body")
    builder.save(docx_path)
    docx_bytes = docx_path.read_bytes()

    def forbidden_read(_path):
        raise AssertionError("parser reread the source path after bytes were captured")

    monkeypatch.setattr(Path, "read_bytes", forbidden_read)
    cases = [
        parse_text(
            tmp_path / "missing.txt",
            raw_bytes=b"captured text",
            logical_path="alpha/shared.txt",
        ),
        parse_md_file(
            tmp_path / "missing.md",
            raw_bytes=b"# Captured\n\nbody",
            logical_path="beta/shared.md",
        ),
        parse_docx_file(
            tmp_path / "missing.docx",
            raw_bytes=docx_bytes,
            logical_path="gamma/shared.docx",
        ),
    ]

    expected = [b"captured text", b"# Captured\n\nbody", docx_bytes]
    for result, captured in zip(cases, expected, strict=True):
        assert result.source.sha256 == hashlib.sha256(captured).hexdigest()
        assert result.source.size == len(captured)
    assert [result.source.logical_path for result in cases] == [
        "alpha/shared.txt",
        "beta/shared.md",
        "gamma/shared.docx",
    ]


def test_stable_logical_path_controls_source_identity(tmp_path):
    first = parse_text(
        tmp_path / "machine-a.txt",
        raw_bytes=b"same",
        logical_path="corpus/same.txt",
    )
    second = parse_text(
        tmp_path / "machine-b.txt",
        raw_bytes=b"same",
        logical_path="corpus/same.txt",
    )
    assert first.source.id == second.source.id
