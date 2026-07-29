import zipfile

from docx import Document as DocxBuilder
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from veriformis.ir import CodeBlock, Heading, ListBlock, Paragraph, Table, block_text
from veriformis.parsers.docx import parse_docx_file as _parse_docx_file


def parse_docx_file(path):
    return _parse_docx_file(path, logical_path=path.name)


def _build(path):
    d = DocxBuilder()
    d.add_heading("Report", level=1)
    d.add_paragraph("Opening paragraph.")
    p = d.add_paragraph()
    p.add_run("Mixed ").bold = False
    run = p.add_run("bold")
    run.bold = True
    p.add_run(" text.")
    d.add_paragraph("First item", style="List Bullet")
    d.add_paragraph("Second item", style="List Bullet")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "a"
    table.cell(1, 1).text = "b"
    d.save(path)


def test_docx_structure_and_provenance(tmp_path):
    path = tmp_path / "sample.docx"
    _build(path)
    result = parse_docx_file(path)
    doc = result.document
    assert result.source.parser == "docx"
    types = [type(b) for b in doc.children]
    assert Heading in types and Paragraph in types
    assert ListBlock in types and Table in types
    heading = next(b for b in doc.children if isinstance(b, Heading))
    assert heading.level == 1 and block_text(heading) == "Report"
    # provenance: every span indexes into the extracted stream
    stream = result.source.extracted_text
    for block in doc.children:
        assert block.span is not None
        assert block.span.start < block.span.end <= len(stream)
    para = next(
        b
        for b in doc.children
        if isinstance(b, Paragraph) and "Opening" in block_text(b)
    )
    assert "Opening paragraph." in stream[para.span.start : para.span.end]
    table = next(block for block in doc.children if isinstance(block, Table))
    assert table.headers == []
    assert len(table.rows) == 2


def test_docx_recovers_body_and_paragraph_wrapper_text_with_diagnostics(tmp_path):
    path = tmp_path / "wrapped.docx"
    builder = DocxBuilder()

    body_paragraph = builder.add_paragraph("Body wrapper text")
    body = builder.element.body
    body_index = list(body).index(body_paragraph._p)
    body.remove(body_paragraph._p)
    body_wrapper = OxmlElement("w:customXml")
    body_wrapper.append(body_paragraph._p)
    body.insert(body_index, body_wrapper)

    paragraph = builder.add_paragraph()
    paragraph.add_run("before ")
    wrapped_run = paragraph.add_run("smart-tag text")
    paragraph.add_run(" after")
    paragraph._p.remove(wrapped_run._r)
    smart_tag = OxmlElement("w:smartTag")
    smart_tag.set(qn("w:uri"), "urn:veriformis:test")
    smart_tag.set(qn("w:element"), "token")
    smart_tag.append(wrapped_run._r)
    paragraph._p.insert(1, smart_tag)
    builder.save(path)

    first = parse_docx_file(path)
    second = parse_docx_file(path)

    assert "Body wrapper text" in first.source.extracted_text
    assert "before smart-tag text after" in first.source.extracted_text
    wrappers = [
        item
        for item in first.diagnostics.diagnostics
        if item.code
        in {
            "docx.body-wrapper-normalized",
            "docx.paragraph-wrapper-normalized",
        }
    ]
    assert {item.code for item in wrappers} == {
        "docx.body-wrapper-normalized",
        "docx.paragraph-wrapper-normalized",
    }
    assert all(item.disposition == "normalized" for item in wrappers)
    assert all(item.loss_kind == "structure" for item in wrappers)
    assert all(
        item.location.part == "word/document.xml" and item.location.xpath
        for item in wrappers
    )
    assert [item.diagnostic_id for item in wrappers] == [
        item.diagnostic_id
        for item in second.diagnostics.diagnostics
        if item.code
        in {
            "docx.body-wrapper-normalized",
            "docx.paragraph-wrapper-normalized",
        }
    ]


def test_docx_reports_unsupported_body_and_run_constructs(tmp_path):
    path = tmp_path / "unsupported.docx"
    builder = DocxBuilder()
    paragraph = builder.add_paragraph()
    run = paragraph.add_run("visible")

    symbol = OxmlElement("w:sym")
    symbol.set(qn("w:font"), "Wingdings")
    symbol.set(qn("w:char"), "F0FC")
    run._r.append(symbol)

    instruction = OxmlElement("w:instrText")
    instruction.text = " PAGE "
    run._r.append(instruction)

    unsupported_run = OxmlElement("w:object")
    hidden_text = OxmlElement("w:t")
    hidden_text.text = "not silently lost"
    unsupported_run.append(hidden_text)
    run._r.append(unsupported_run)

    unsupported_body = OxmlElement("w:altChunk")
    unsupported_body.set(qn("r:id"), "rIdUnsupported")
    body = builder.element.body
    body.insert(len(body) - 1, unsupported_body)
    builder.save(path)

    result = parse_docx_file(path)
    assert result.source.extracted_text == "visible"

    by_code = {item.code: item for item in result.diagnostics.diagnostics}
    expected = {
        "docx.unsupported-body-element",
        "docx.symbol-omitted",
        "docx.field-instruction-omitted",
        "docx.unsupported-run-element",
    }
    assert expected <= set(by_code)
    assert by_code["docx.unsupported-body-element"].loss_kind == "text"
    assert by_code["docx.symbol-omitted"].details == {
        "element": "sym",
        "font": "Wingdings",
        "char": "F0FC",
    }
    assert by_code["docx.field-instruction-omitted"].loss_kind == "metadata"
    assert by_code["docx.unsupported-run-element"].details["contains_text"] is True
    for code in expected:
        item = by_code[code]
        assert item.disposition == "omitted"
        assert item.location.kind == "ooxml"
        assert item.location.part == "word/document.xml"
        assert item.location.xpath


def test_docx_recovers_supported_non_text_run_constructs(tmp_path):
    path = tmp_path / "run-constructs.docx"
    builder = DocxBuilder()
    run = builder.add_paragraph().add_run("a")
    run._r.append(OxmlElement("w:cr"))
    tail = OxmlElement("w:t")
    tail.text = "b"
    run._r.append(tail)
    run._r.append(OxmlElement("w:noBreakHyphen"))
    final = OxmlElement("w:t")
    final.text = "c"
    run._r.append(final)
    builder.save(path)

    result = parse_docx_file(path)
    assert result.source.extracted_text == "a\nb‑c"
    assert not any(
        item.code == "docx.unsupported-run-element"
        for item in result.diagnostics.diagnostics
    )


def test_docx_consumed_core_elements_do_not_get_unsupported_diagnostics(tmp_path):
    path = tmp_path / "supported.docx"
    builder = DocxBuilder()
    paragraph = builder.add_paragraph("plain ")
    paragraph.add_run("bold").bold = True
    builder.save(path)

    result = parse_docx_file(path)
    assert not any(
        item.code.startswith("docx.unsupported-")
        for item in result.diagnostics.diagnostics
    )


def test_docx_note_blocks_have_canonical_stream_spans(tmp_path):
    path = tmp_path / "notes.docx"
    builder = DocxBuilder()
    run = builder.add_paragraph("Body text").add_run()
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "1")
    run._r.append(reference)
    builder.save(path)

    word_namespace = qn("w:footnotes").split("}")[0][1:]
    footnotes_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="{word_namespace}">
  <w:footnote w:id="1">
    <w:p><w:r><w:t>Footnote text</w:t></w:r></w:p>
  </w:footnote>
</w:footnotes>
""".encode()
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr("word/footnotes.xml", footnotes_xml)

    result = parse_docx_file(path)
    body_block = result.document.children[0]
    note_block = result.document.footnotes["1"].children[0]

    assert result.source.extracted_text == "Body text[^1]\n\nFootnote text"
    assert body_block.block_index == 0
    assert note_block.block_index == 1
    assert body_block.span is not None
    assert note_block.span is not None
    assert (
        result.source.extracted_text[body_block.span.start : body_block.span.end]
        == "Body text[^1]"
    )
    assert (
        result.source.extracted_text[note_block.span.start : note_block.span.end]
        == "Footnote text"
    )


def test_docx_note_part_loss_is_located_in_the_note_part(tmp_path):
    path = tmp_path / "note-loss.docx"
    builder = DocxBuilder()
    run = builder.add_paragraph("Body").add_run()
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "1")
    run._r.append(reference)
    builder.save(path)

    word_namespace = qn("w:footnotes").split("}")[0][1:]
    footnotes_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<w:footnotes xmlns:w="{word_namespace}">
  <w:footnote w:id="1">
    <w:p><w:r><w:t>Visible note</w:t><w:sym w:font="Wingdings" w:char="F0FC"/></w:r></w:p>
    <w:del><w:r><w:delText>Deleted note text</w:delText></w:r></w:del>
  </w:footnote>
</w:footnotes>
""".encode()
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr("word/footnotes.xml", footnotes_xml)

    result = parse_docx_file(path)
    matching = [
        item
        for item in result.diagnostics.diagnostics
        if item.code in {"docx.symbol-omitted", "docx.revision-deletion-omitted"}
    ]

    assert {item.code for item in matching} == {
        "docx.symbol-omitted",
        "docx.revision-deletion-omitted",
    }
    assert all(item.location.part == "word/footnotes.xml" for item in matching)
    assert all(item.location.xpath for item in matching)
    assert block_text(result.document.footnotes["1"].children[0]) == "Visible note"


def test_docx_malformed_note_part_refuses_canonical_recovery(tmp_path):
    path = tmp_path / "malformed-note.docx"
    builder = DocxBuilder()
    builder.add_paragraph("Body")
    builder.save(path)
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr("word/footnotes.xml", b"<w:footnotes")

    result = parse_docx_file(path)
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "docx.footnote-part-invalid"
    )

    assert result.diagnostics.status == "refused"
    assert item.disposition == "refused"
    assert item.location.part == "word/footnotes.xml"


def test_docx_unresolved_note_reference_refuses_canonical_recovery(tmp_path):
    path = tmp_path / "unresolved-note.docx"
    builder = DocxBuilder()
    run = builder.add_paragraph("Body").add_run()
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "404")
    run._r.append(reference)
    builder.save(path)

    result = parse_docx_file(path)
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "docx.footnote-reference-unresolved"
    )

    assert result.diagnostics.status == "refused"
    assert item.details["note_id"] == "404"
    assert item.location.part == "word/document.xml"


def test_docx_uses_only_explicitly_marked_table_header_rows(tmp_path):
    path = tmp_path / "table-header.docx"
    builder = DocxBuilder()
    table = builder.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Header"
    table.cell(1, 0).text = "Value"
    row_properties = OxmlElement("w:trPr")
    row_properties.append(OxmlElement("w:tblHeader"))
    table.rows[0]._tr.insert(0, row_properties)
    builder.save(path)

    parsed = next(
        block
        for block in parse_docx_file(path).document.children
        if isinstance(block, Table)
    )

    assert [cell.children[0].value for cell in parsed.headers] == ["Header"]
    assert [[cell.children[0].value for cell in row] for row in parsed.rows] == [
        ["Value"]
    ]


def test_docx_refuses_header_rows_beyond_canonical_single_header(tmp_path):
    path = tmp_path / "multiple-table-headers.docx"
    builder = DocxBuilder()
    table = builder.add_table(rows=2, cols=1)
    for row in table.rows:
        properties = OxmlElement("w:trPr")
        properties.append(OxmlElement("w:tblHeader"))
        row._tr.insert(0, properties)
    builder.save(path)

    result = parse_docx_file(path)
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "docx.table-header-row-unsupported"
    )

    assert result.diagnostics.status == "refused"
    assert item.details["row_index"] == 1


def test_docx_nested_table_text_and_merged_cells_are_refused(tmp_path):
    nested_path = tmp_path / "nested-table.docx"
    nested_builder = DocxBuilder()
    outer = nested_builder.add_table(rows=1, cols=1)
    outer.cell(0, 0).add_table(rows=1, cols=1).cell(0, 0).text = "Nested text"
    nested_builder.save(nested_path)

    nested = parse_docx_file(nested_path)
    nested_item = next(
        item
        for item in nested.diagnostics.diagnostics
        if item.code == "docx.nested-table-unsupported"
    )
    assert nested.diagnostics.status == "refused"
    assert nested_item.loss_kind == "text"
    assert nested_item.location.xpath

    merged_path = tmp_path / "merged-table.docx"
    merged_builder = DocxBuilder()
    merged = merged_builder.add_table(rows=1, cols=2)
    merged.cell(0, 0).text = "Left"
    merged.cell(0, 1).text = "Right"
    merged.cell(0, 0).merge(merged.cell(0, 1))
    merged_builder.save(merged_path)

    merged_result = parse_docx_file(merged_path)
    merge_item = next(
        item
        for item in merged_result.diagnostics.diagnostics
        if item.code == "docx.table-cell-merge-unsupported"
    )
    assert merged_result.diagnostics.status == "refused"
    assert merge_item.disposition == "refused"
    assert merge_item.details["merge_kind"] == "gridSpan"


def test_docx_source_code_preserves_all_supported_text_controls(tmp_path):
    path = tmp_path / "source-code.docx"
    builder = DocxBuilder()
    paragraph = builder.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "SourceCode")
    properties.insert(0, style)
    run = paragraph.add_run("a")
    run._r.append(OxmlElement("w:cr"))
    middle = OxmlElement("w:t")
    middle.text = "b"
    run._r.append(middle)
    run._r.append(OxmlElement("w:noBreakHyphen"))
    tail = OxmlElement("w:t")
    tail.text = "c"
    run._r.append(tail)
    run._r.append(OxmlElement("w:softHyphen"))
    final = OxmlElement("w:t")
    final.text = "d"
    run._r.append(final)
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run._r.append(page_break)
    unsupported = OxmlElement("w:object")
    hidden = OxmlElement("w:t")
    hidden.text = "hidden"
    unsupported.append(hidden)
    run._r.append(unsupported)
    builder.save(path)

    result = parse_docx_file(path)
    code = next(
        block for block in result.document.children if isinstance(block, CodeBlock)
    )

    assert code.text == "a\nb‑c\u00add"
    assert {
        "docx.page-break-omitted",
        "docx.unsupported-run-element",
    } <= {item.code for item in result.diagnostics.diagnostics}


def test_docx_note_marker_removal_preserves_payload_in_same_run(tmp_path):
    path = tmp_path / "marker-payload.docx"
    builder = DocxBuilder()
    reference_run = builder.add_paragraph("Body").add_run()
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "1")
    reference_run._r.append(reference)
    builder.save(path)

    word_namespace = qn("w:footnotes").split("}")[0][1:]
    footnotes_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<w:footnotes xmlns:w="{word_namespace}">
  <w:footnote w:id="1">
    <w:p><w:r><w:footnoteRef/><w:t>Payload beside marker</w:t></w:r></w:p>
  </w:footnote>
</w:footnotes>
""".encode()
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr("word/footnotes.xml", footnotes_xml)

    result = parse_docx_file(path)

    assert (
        block_text(result.document.footnotes["1"].children[0])
        == "Payload beside marker"
    )


def test_docx_duplicate_note_ids_and_continuation_notice_are_inventoried(tmp_path):
    path = tmp_path / "duplicate-notes.docx"
    builder = DocxBuilder()
    reference_run = builder.add_paragraph("Body").add_run()
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "1")
    reference_run._r.append(reference)
    builder.save(path)

    word_namespace = qn("w:footnotes").split("}")[0][1:]
    footnotes_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<w:footnotes xmlns:w="{word_namespace}">
  <w:footnote w:type="continuationNotice" w:id="-3">
    <w:p><w:r><w:t>Generated notice</w:t></w:r></w:p>
  </w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:t>First body</w:t></w:r></w:p></w:footnote>
  <w:footnote w:id="1"><w:p><w:r><w:t>Second body</w:t></w:r></w:p></w:footnote>
</w:footnotes>
""".encode()
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr("word/footnotes.xml", footnotes_xml)

    result = parse_docx_file(path)
    duplicate = next(
        item
        for item in result.diagnostics.diagnostics
        if item.code == "docx.footnote-id-duplicate"
    )
    notice = next(
        item
        for item in result.diagnostics.diagnostics
        if item.code == "docx.footnote-separator-omitted"
        and item.details["note_type"] == "continuationNotice"
    )

    assert result.diagnostics.status == "refused"
    assert block_text(result.document.footnotes["1"].children[0]) == "First body"
    assert duplicate.location.part == "word/footnotes.xml"
    assert notice.disposition == "omitted"
