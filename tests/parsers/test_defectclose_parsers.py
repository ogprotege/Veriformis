"""Defect-closure regressions: html/docx/markdown recovery loss must be explicit.

Covers four verified defects:

1. HTML silently discarded visible text outside the handled block-tag set
   (text directly in ``<div>``/``<span>``/``<body>``) and destroyed tail text
   following captured elements, while reporting status ``complete``.
2. Markdown footnote definitions wrapped in blockquotes or list items were
   invisible to the definition pre-pass, so unused and duplicate definitions
   inside containers vanished or replaced note text without any diagnostic.
3. DOCX table rows, cells, and cell paragraphs wrapped in ``w:sdt`` /
   ``w:customXml`` (Word content controls) silently vanished from tables.
4. DOCX code-classified runs erased line breaks and note-reference markers
   while the run-boundary inventory marked them handled.
"""

from __future__ import annotations

import zipfile

from docx import Document as DocxBuilder
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from veriformis.ir import (
    Code,
    FootnoteRef,
    Heading,
    Paragraph,
    Table,
    block_text,
    validate_document_against_stream,
)
from veriformis.parsers.docx import parse_docx_file as _parse_docx_file
from veriformis.parsers.html import parse_html_file as _parse_html_file
from veriformis.parsers.markdown import parse_md_file as _parse_md_file


def parse_html(path):
    return _parse_html_file(path, logical_path=path.name)


def parse_md(path):
    return _parse_md_file(path, logical_path=path.name)


def parse_docx(path):
    return _parse_docx_file(path, logical_path=path.name)


def _assert_exact(result) -> None:
    validate_document_against_stream(
        result.document,
        result.source.extracted_text,
        exact=True,
    )


# ---------------------------------------------------------------------------
# Defect 1: HTML loose text outside the handled block-tag set
# ---------------------------------------------------------------------------


def test_html_recovers_div_text_after_captured_block(tmp_path):
    path = tmp_path / "div-loss.html"
    path.write_bytes(
        b"<html><body><h1>Title</h1><div>Loose text in a div</div></body></html>"
    )

    result = parse_html(path)

    assert result.source.extracted_text == "Title\n\nLoose text in a div"
    types = [type(block) for block in result.document.children]
    assert types == [Heading, Paragraph]
    _assert_exact(result)
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "html.loose-text-recovered"
    )
    assert item.disposition == "normalized"
    assert result.diagnostics.status == "degraded"


def test_html_recovers_tail_text_after_captured_elements_in_order(tmp_path):
    path = tmp_path / "tail-loss.html"
    path.write_bytes(b"<html><body><h1>T</h1>tail text here<p>P</p></body></html>")

    result = parse_html(path)

    assert result.source.extracted_text == "T\n\ntail text here\n\nP"
    _assert_exact(result)


def test_html_recovers_direct_body_text_in_document_order(tmp_path):
    path = tmp_path / "body-loss.html"
    path.write_bytes(b"<html><body>direct body text<p>P</p></body></html>")

    result = parse_html(path)

    assert result.source.extracted_text == "direct body text\n\nP"
    _assert_exact(result)


def test_html_recovers_span_text_between_captured_blocks(tmp_path):
    path = tmp_path / "span-loss.html"
    path.write_bytes(
        b"<html><body><p>P</p><span>span only text</span></body></html>"
    )

    result = parse_html(path)

    assert result.source.extracted_text == "P\n\nspan only text"
    _assert_exact(result)


def test_html_loose_text_recovery_is_deterministic(tmp_path):
    path = tmp_path / "deterministic.html"
    path.write_bytes(
        b"<html><body><h1>T</h1>tail<div>div text</div><p>P</p></body></html>"
    )

    first = parse_html(path)
    second = parse_html(path)

    assert first.source.extracted_text == second.source.extracted_text
    assert first.diagnostics.report_digest == second.diagnostics.report_digest


# ---------------------------------------------------------------------------
# Defect 2: Markdown container-wrapped footnote definitions
# ---------------------------------------------------------------------------


def test_markdown_unused_definition_in_blockquote_is_refused(tmp_path):
    path = tmp_path / "quote-unused.md"
    path.write_text(
        "Intro paragraph.\n\n> [^lost]: This note body would vanish.\n",
        encoding="utf-8",
    )

    result = parse_md(path)
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "markdown.unused-footnote-definition-refused"
    )

    assert result.diagnostics.status == "refused"
    assert item.loss_kind == "text"
    assert item.location.line_start == 3
    assert item.details == {"label": "lost"}


def test_markdown_duplicate_definition_in_blockquote_is_refused(tmp_path):
    path = tmp_path / "quote-duplicate.md"
    path.write_text(
        "Use[^n].\n\n[^n]: first body\n\n> [^n]: second body inside quote\n",
        encoding="utf-8",
    )

    result = parse_md(path)
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "markdown.duplicate-footnote-definition"
    )

    assert result.diagnostics.status == "refused"
    assert item.location.line_start == 5
    assert item.details == {"label": "n", "first_definition_line": 3}


def test_markdown_unused_definition_in_list_item_is_refused(tmp_path):
    # The mdit footnote plugin recognizes definitions inside list items
    # (verified empirically), so the pre-pass must inventory them too.
    path = tmp_path / "list-unused.md"
    path.write_text(
        "Body paragraph.\n\n- [^li]: list-wrapped note body\n",
        encoding="utf-8",
    )

    result = parse_md(path)
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "markdown.unused-footnote-definition-refused"
    )

    assert result.diagnostics.status == "refused"
    assert item.location.line_start == 3
    assert item.details == {"label": "li"}


def test_markdown_referenced_container_definitions_stay_clean(tmp_path):
    path = tmp_path / "container-used.md"
    path.write_text(
        "Use[^q] and[^li].\n\n> [^q]: quoted body\n\n- [^li]: listed body\n",
        encoding="utf-8",
    )

    result = parse_md(path)

    assert set(result.document.footnotes) == {"q", "li"}
    assert not any(
        item.code
        in {
            "markdown.unused-footnote-definition-refused",
            "markdown.duplicate-footnote-definition",
        }
        for item in result.diagnostics.diagnostics
    )


# ---------------------------------------------------------------------------
# Defect 3: DOCX sdt/customXml-wrapped table rows, cells, cell paragraphs
# ---------------------------------------------------------------------------


def _wrap_in_sdt(element):
    wrapper = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    wrapper.append(content)
    return wrapper, content


def test_docx_recovers_sdt_wrapped_cell_paragraphs(tmp_path):
    path = tmp_path / "sdt-cell.docx"
    builder = DocxBuilder()
    table = builder.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "plain cell"
    table.cell(0, 1).text = "wrapped cell text"
    cell = table.cell(0, 1)._tc
    wrapper, content = _wrap_in_sdt(cell)
    for paragraph in cell.findall(qn("w:p")):
        cell.remove(paragraph)
        content.append(paragraph)
    cell.append(wrapper)
    builder.save(path)

    result = parse_docx(path)
    parsed = next(
        block for block in result.document.children if isinstance(block, Table)
    )

    assert "wrapped cell text" in result.source.extracted_text
    assert block_text(parsed) == "plain cell\twrapped cell text"
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "docx.table-wrapper-normalized"
    )
    assert item.disposition == "normalized"
    assert item.location.part == "word/document.xml"
    assert item.location.xpath
    _assert_exact(result)


def test_docx_recovers_sdt_wrapped_table_row_and_cell(tmp_path):
    path = tmp_path / "sdt-row.docx"
    builder = DocxBuilder()
    table = builder.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "first row"
    table.cell(1, 0).text = "wrapped row"
    row = table.rows[1]._tr
    tbl = table._tbl
    index = list(tbl).index(row)
    tbl.remove(row)
    row_wrapper, row_content = _wrap_in_sdt(row)
    tbl.insert(index, row_wrapper)

    cell = row.find(qn("w:tc"))
    row.remove(cell)
    cell_wrapper, cell_content = _wrap_in_sdt(cell)
    cell_content.append(cell)
    row.append(cell_wrapper)
    row_content.append(row)
    builder.save(path)

    result = parse_docx(path)
    parsed = next(
        block for block in result.document.children if isinstance(block, Table)
    )

    assert block_text(parsed) == "first row\nwrapped row"
    assert any(
        diagnostic.code == "docx.table-wrapper-normalized"
        for diagnostic in result.diagnostics.diagnostics
    )
    _assert_exact(result)


# ---------------------------------------------------------------------------
# Defect 4: DOCX code-classified runs erase breaks and note references
# ---------------------------------------------------------------------------


def _add_footnotes_part(path, note_text="Note body"):
    word_namespace = qn("w:footnotes").split("}")[0][1:]
    footnotes_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="{word_namespace}">
  <w:footnote w:id="1">
    <w:p><w:r><w:t>{note_text}</w:t></w:r></w:p>
  </w:footnote>
</w:footnotes>
""".encode()
    with zipfile.ZipFile(path, "a") as archive:
        archive.writestr("word/footnotes.xml", footnotes_xml)


def test_docx_code_run_preserves_line_breaks(tmp_path):
    path = tmp_path / "code-break.docx"
    builder = DocxBuilder()
    run = builder.add_paragraph().add_run("line1")
    run.font.name = "Consolas"
    run._r.append(OxmlElement("w:br"))
    tail = OxmlElement("w:t")
    tail.text = "line2"
    run._r.append(tail)
    builder.save(path)

    result = parse_docx(path)
    paragraph = result.document.children[0]
    code = next(node for node in paragraph.children if isinstance(node, Code))

    assert code.value == "line1\nline2"
    assert "line1\nline2" in result.source.extracted_text
    _assert_exact(result)


def test_docx_code_run_preserves_footnote_reference_marker(tmp_path):
    path = tmp_path / "code-note-ref.docx"
    builder = DocxBuilder()
    run = builder.add_paragraph().add_run("code")
    run.font.name = "Consolas"
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "1")
    run._r.append(reference)
    builder.save(path)
    _add_footnotes_part(path)

    result = parse_docx(path)
    paragraph = result.document.children[0]

    assert [type(node) for node in paragraph.children] == [Code, FootnoteRef]
    assert paragraph.children[0].value == "code"
    assert paragraph.children[1].id == "1"
    assert result.source.extracted_text == "code[^1]\n\nNote body"
    _assert_exact(result)


def test_docx_code_block_paragraph_note_reference_loss_is_located(tmp_path):
    path = tmp_path / "code-block-note-ref.docx"
    builder = DocxBuilder()
    paragraph = builder.add_paragraph()
    properties = paragraph._p.get_or_add_pPr()
    style = OxmlElement("w:pStyle")
    style.set(qn("w:val"), "SourceCode")
    properties.insert(0, style)
    run = paragraph.add_run("block code")
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "1")
    run._r.append(reference)
    builder.save(path)
    _add_footnotes_part(path)

    result = parse_docx(path)
    item = next(
        diagnostic
        for diagnostic in result.diagnostics.diagnostics
        if diagnostic.code == "docx.code-block-inline-omitted"
    )

    assert item.disposition == "omitted"
    assert item.location.part == "word/document.xml"
    assert item.location.xpath
    assert result.source.extracted_text == "block code\n\nNote body"


def test_docx_code_run_recovery_is_deterministic(tmp_path):
    path = tmp_path / "code-deterministic.docx"
    builder = DocxBuilder()
    run = builder.add_paragraph().add_run("a")
    run.font.name = "Consolas"
    run._r.append(OxmlElement("w:br"))
    tail = OxmlElement("w:t")
    tail.text = "b"
    run._r.append(tail)
    builder.save(path)

    first = parse_docx(path)
    second = parse_docx(path)

    assert first.source.extracted_text == second.source.extracted_text
    assert first.diagnostics.report_digest == second.diagnostics.report_digest
